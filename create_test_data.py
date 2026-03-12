from docx import Document

def create_resume():
    doc = Document()
    doc.add_heading('John Doe', 0)
    doc.add_paragraph('Software Engineer with 5 years of experience in Python, FastAPI, and Machine Learning.')
    
    doc.add_heading('Skills', level=1)
    doc.add_paragraph('Python, FastAPI, SQL, Docker, AWS, Scikit-learn, NLP')
    
    doc.add_heading('Experience', level=1)
    doc.add_paragraph('Senior Developer at Tech Corp (2019-Present)')
    doc.add_paragraph('- Developed high-performance APIs using FastAPI.')
    doc.add_paragraph('- Implemented machine learning models for data analysis.')
    
    doc.save('test_resume.docx')
    print("test_resume.docx created successfully.")

if __name__ == "__main__":
    create_resume()
